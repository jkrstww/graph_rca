import chardet
import re
import os
from typing import Optional

def get_encoding(file_path):
    """
    检测文件的编码格式，并正确处理文件路径

    参数:
        file_path: 文件路径，可以是绝对路径或相对路径

    返回:
        文件的编码格式
    """

    # 如果传入的是相对路径，尝试多种方式解析
    if not os.path.isabs(file_path):
        # 方法1: 基于当前工作目录
        abs_path_from_cwd = os.path.abspath(file_path)

        # 方法2: 基于调用者位置
        try:
            import inspect
            caller_frame = inspect.stack()[1]
            caller_file = caller_frame.filename
            caller_dir = os.path.dirname(os.path.abspath(caller_file))
            abs_path_from_caller = os.path.join(caller_dir, file_path)
        except (IndexError, AttributeError):
            abs_path_from_caller = None

        # 方法3: 基于当前脚本位置
        current_script_dir = os.path.dirname(os.path.abspath(__file__))
        abs_path_from_script = os.path.join(current_script_dir, file_path)

        # 尝试所有可能的路径，选择第一个存在的
        possible_paths = [
            abs_path_from_cwd,
            abs_path_from_caller,
            abs_path_from_script
        ]

        # 移除None值
        possible_paths = [p for p in possible_paths if p is not None]

        # 尝试每个路径
        for path in possible_paths:
            normalized_path = os.path.normpath(path)
            if os.path.exists(normalized_path):
                file_path = normalized_path
                break
        else:
            # 如果没有找到任何存在的路径，使用基于当前工作目录的路径
            file_path = os.path.normpath(abs_path_from_cwd)

    # 规范化路径（处理 ../ 和 ./ 等）
    file_path = os.path.normpath(file_path)

    # 检查文件是否存在
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")

    if not os.path.isfile(file_path):
        raise IsADirectoryError(f"路径指向的是目录而不是文件: {file_path}")

    with open(file_path, 'rb') as f:
        tmp = chardet.detect(f.read())
        return tmp['encoding']

def extract_label_content(text: str, label: str) -> list:
    """
    从文本中提取所有被<label></label>标签包含的内容

    Args:
        text: 要搜索的文本
        label: 标签名

    Returns:
        包含所有匹配内容的列表，如果没有找到则返回空列表
    """
    # 构建正则表达式模式，注意转义标签名
    pattern = f"<{label}>(.*?)</{label}>"

    # 使用非贪婪模式查找所有匹配项
    matches = re.findall(pattern, text, re.DOTALL)

    # 返回匹配结果列表
    return matches if matches else []

def read_document_content(file_path: str) -> str:
    """
    从txt、doc、docx、pdf文件中读取文本内容
    
    参数:
        file_path: 文件路径
        
    返回:
        str: 提取的文本内容，如果读取失败则返回None
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")
    
    file_ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_ext == '.txt':
            return _read_txt(file_path)
        elif file_ext == '.docx':
            return _read_docx(file_path)
        elif file_ext == '.doc':
            return _read_doc(file_path)
        elif file_ext == '.pdf':
            return _read_pdf(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {file_ext}")
    except Exception as e:
        print(f"读取文件时出错 {file_path}: {str(e)}")
        return ''

def _read_txt(file_path: str) -> str:
    """读取txt文件内容"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # 如果utf-8失败，尝试其他编码
        with open(file_path, 'r', encoding='gbk') as file:
            return file.read()

def _read_docx(file_path: str) -> str:
    """读取docx文件内容"""
    try:
        from docx import Document
        doc = Document(file_path)
        text_content = []
        for paragraph in doc.paragraphs:
            text_content.append(paragraph.text)
        
        # 读取表格内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_content.append(cell.text)
        
        return '\n'.join(text_content)
    except ImportError:
        raise ImportError("请安装python-docx库: pip install python-docx")

def _read_doc(file_path: str) -> str:
    """读取doc文件内容"""
    try:
        # 方法1: 使用antiword（需要安装antiword工具）
        import subprocess
        try:
            result = subprocess.run(['antiword', file_path], 
                                  capture_output=True, text=True, timeout=30)
            if result.returncode == 0:
                return result.stdout
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass
        
        # 方法2: 使用python-docx2txt（备用方案）
        try:
            import docx2txt
            return docx2txt.process(file_path)
        except ImportError:
            pass
        
        # 方法3: 使用win32com（仅Windows系统）
        # try:
        #     import win32com.client
        #     word = win32com.client.Dispatch("Word.Application")
        #     word.visible = False
        #     doc = word.Documents.Open(file_path)
        #     text = doc.Content.Text
        #     doc.Close()
        #     word.Quit()
        #     return text
        # except ImportError:
        #     pass
        
        raise ImportError("无法读取doc文件，请安装以下任一工具:\n"
                         "1. antiword (Linux/Mac): sudo apt-get install antiword\n"
                         "2. python-docx2txt: pip install docx2txt\n"
                         "3. pywin32 (Windows): pip install pywin32")
                         
    except Exception as e:
        raise Exception(f"读取doc文件失败: {str(e)}")

def _read_pdf(file_path: str) -> str:
    """读取pdf文件内容"""
    try:
        # 方法1: 使用PyPDF2
        try:
            import PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = []
                for page in pdf_reader.pages:
                    text_content.append(page.extract_text())
                return '\n'.join(text_content)
        except ImportError:
            pass
        
        # 方法2: 使用pdfplumber（更准确）
        try:
            import pdfplumber
            text_content = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
            return '\n'.join(text_content)
        except ImportError:
            pass
        
        # 方法3: 使用pdfminer
        try:
            from pdfminer.high_level import extract_text
            return extract_text(file_path)
        except ImportError:
            pass
        
        raise ImportError("无法读取pdf文件，请安装以下任一库:\n"
                         "1. PyPDF2: pip install PyPDF2\n"
                         "2. pdfplumber: pip install pdfplumber\n"
                         "3. pdfminer.six: pip install pdfminer.six")
                         
    except Exception as e:
        raise Exception(f"读取pdf文件失败: {str(e)}")

def clean_text(text: str) -> str:
    """
    清理文本内容，去除多余的空白字符
    
    参数:
        text: 原始文本
        
    返回:
        str: 清理后的文本
    """
    if not text:
        return ""
    
    # 替换多个连续空白字符为单个空格
    text = re.sub(r'\s+', ' ', text)
    # 去除首尾空白
    text = text.strip()
    
    return text

if __name__ == '__main__':
    sample_text = """
    <events>
    <event>变压器温度升高</event>
    <event>绝缘油泄漏</event>
    <event>保护装置动作</event>
    </events>
    还有一些其他文本<event>额外事件</event>
    """

    # 提取event标签内容
    events = extract_label_content(sample_text, "event")
    print(events)  # 输出: ['变压器温度升高', '绝缘油泄漏', '保护装置动作', '额外事件']

    # 提取不存在的标签内容
    none_results = extract_label_content(sample_text, "nonexistent")
    print(none_results)  # 输出: []